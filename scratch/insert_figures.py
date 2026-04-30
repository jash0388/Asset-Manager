#!/usr/bin/env python3
"""
Insert Figures 6.1, 6.2, 6.3 into the AlphaBayX report document.
"""

import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# --- Paths ---
INPUT_DOC  = "/Users/jashwanthsingh/Downloads/AlphaBayX_RTRP_Report_fixed-2.docx"
OUTPUT_DOC = "/Users/jashwanthsingh/Downloads/AlphaBayX_RTRP_Report_Final.docx"
IMG_DIR    = "/Users/jashwanthsingh/.gemini/antigravity/brain/c7bd3b9d-5889-4b62-92fe-1a014609bccf"

FIG_6_1 = os.path.join(IMG_DIR, "figure_6_1_home_feed_1777483172214.png")
FIG_6_2 = os.path.join(IMG_DIR, "figure_6_2_admin_dashboard_1777483193238.png")
FIG_6_3 = os.path.join(IMG_DIR, "figure_6_3_exam_portal.png")

# Map search text → (image_path, caption_text)
FIGURE_TARGETS = {
    "Figure 6.1": (FIG_6_1, "Figure 6.1: Result Walkthrough — Home Feed (datanauts.in)"),
    "Figure 6.2": (FIG_6_2, "Figure 6.2: Result Walkthrough — Admin Dashboard"),
    "Figure 6.3": (FIG_6_3, "Figure 6.3: Result Walkthrough — Exam Portal (sphn.online)"),
}

def add_figure_after_para(doc, para_idx, image_path, caption):
    """Insert an image + caption paragraph after para_idx."""
    # Add image paragraph
    img_para = OxmlElement('w:p')
    
    # We'll insert by manipulating the body XML
    body = doc.element.body
    ref_para = doc.paragraphs[para_idx]._element
    
    # Create image paragraph
    new_doc_para = doc.add_paragraph()
    new_doc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_doc_para.add_run()
    run.add_picture(image_path, width=Inches(5.5))
    
    # Move it after the reference paragraph
    body.remove(new_doc_para._element)
    ref_para.addnext(new_doc_para._element)
    
    # Create caption paragraph
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(10)
    cap_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    
    # Move caption after image
    body.remove(cap_para._element)
    new_doc_para._element.addnext(cap_para._element)
    
    return cap_para._element  # return caption element for next insertion reference

def main():
    print(f"Loading: {INPUT_DOC}")
    doc = Document(INPUT_DOC)
    
    # Find the paragraphs that mention our figure references
    # We look for the section "6.4 Results Walkthrough" or explicit figure mentions
    target_phrases = {
        "Figures 6.1, 6.2, and 6.3": ["Figure 6.1", "Figure 6.2", "Figure 6.3"],
        "Figures 6.1": ["Figure 6.1"],
        "Figure 6.1 and Figure 6.2": ["Figure 6.1", "Figure 6.2"],
    }
    
    # Find paragraph index containing the results walkthrough section
    results_para_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "6.4" in para.text and "Results" in para.text:
            results_para_idx = i
            print(f"Found Results section at paragraph {i}: {para.text[:80]}")
            break
        if "representative sequence of screens" in para.text.lower():
            results_para_idx = i
            print(f"Found results reference at paragraph {i}: {para.text[:80]}")
            break
    
    if results_para_idx is None:
        # Search broadly
        for i, para in enumerate(doc.paragraphs):
            if "datanauts.in" in para.text and "home feed" in para.text.lower():
                results_para_idx = i
                break
    
    if results_para_idx is None:
        # Just find the paragraph with "Figures 6.1"
        for i, para in enumerate(doc.paragraphs):
            text = para.text
            if "Figure 6.1" in text or "6.1, 6.2" in text:
                results_para_idx = i
                print(f"Found figure reference at para {i}: {text[:80]}")
                break
    
    if results_para_idx is None:
        print("Could not find target section. Will append figures at end.")
        results_para_idx = len(doc.paragraphs) - 3  # Near end
    
    print(f"\nWill insert figures after paragraph {results_para_idx}")
    print(f"Paragraph text: {doc.paragraphs[results_para_idx].text[:100]}")
    
    # Check images exist
    for fig_key, (img_path, caption) in FIGURE_TARGETS.items():
        if os.path.exists(img_path):
            print(f"✓ {fig_key}: {img_path}")
        else:
            print(f"✗ MISSING: {fig_key}: {img_path}")
    
    # Insert figures after the results walkthrough paragraph
    # We insert in reverse order so each one ends up in correct position
    body = doc.element.body
    ref_element = doc.paragraphs[results_para_idx]._element
    
    # Insert 6.3 first (will be pushed down by 6.1 and 6.2)
    for fig_key in ["Figure 6.3", "Figure 6.2", "Figure 6.1"]:
        img_path, caption = FIGURE_TARGETS[fig_key]
        if not os.path.exists(img_path):
            print(f"Skipping {fig_key} - image not found")
            continue
        
        # Spacer paragraph
        spacer = OxmlElement('w:p')
        ref_element.addnext(spacer)
        
        # Caption paragraph
        cap_para = OxmlElement('w:p')
        # ...build with python-docx then move
        tmp_cap = doc.add_paragraph()
        tmp_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = tmp_cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(10)
        body.remove(tmp_cap._element)
        spacer.addnext(tmp_cap._element)
        
        # Image paragraph
        tmp_img = doc.add_paragraph()
        tmp_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tmp_img.add_run()
        try:
            run.add_picture(img_path, width=Inches(5.5))
            body.remove(tmp_img._element)
            spacer.addnext(tmp_img._element)
            print(f"✓ Inserted {fig_key}")
        except Exception as e:
            print(f"✗ Error inserting {fig_key}: {e}")
    
    doc.save(OUTPUT_DOC)
    print(f"\n✓ Document saved: {OUTPUT_DOC}")

if __name__ == "__main__":
    main()
